import tkinter as tk
from tkinter import messagebox
from bs4 import BeautifulSoup
from docx import Document
import requests

# Function to scrape relevant data and equations from a website based on a topic keyword
def scrape_website(topic_keyword, website_url):
    try:
        # Send a GET request to the website
        response = requests.get(website_url)
        response.raise_for_status()

        # Parse the HTML content
        soup = BeautifulSoup(response.text, 'html.parser')

        # Find and extract relevant information based on the topic keyword
        relevant_content = []

        # Iterate through all the paragraphs on the page and check for the keyword
        for paragraph in soup.find_all("p"):
            text = paragraph.get_text()
            if topic_keyword.lower() in text.lower():
                relevant_content.append(text)

        # Find and extract equations (if available)
        equations = []
        for equation in soup.find_all("span", class_="math-tex"):
            equations.append(equation.get_text())

        # Combine the relevant paragraphs and equations into a single string
        content = "\n\n".join(relevant_content + equations)

        # Define keywords to filter out specific lines (customize as needed)
        exclude_keywords = ["Learn more about", "Visit", "Read more at"]

        # Filter out lines containing exclude_keywords
        filtered_lines = []
        for line in content.split("\n"):
            if not any(keyword in line for keyword in exclude_keywords):
                filtered_lines.append(line)

        # Combine the filtered lines into a final string
        filtered_content = "\n".join(filtered_lines)

        return filtered_content

    except Exception as e:
        print(f"Error: {str(e)}")
        return None

# Function to create a Word document with formatted content using python-docx
def create_word_document(topic, content):
    doc = Document()

    # Add the topic as a heading
    doc.add_heading(topic, level=1)

    # Add the scraped content with proper formatting
    for paragraph in content.split("\n"):
        doc.add_paragraph(paragraph)

    # Save the Word document in the root folder
    file_name = f"{topic}_data.docx"
    doc.save(file_name)

    messagebox.showinfo("Success", f"Document for '{topic}' generated successfully.")


# Event handler function for the "Generate Document" button
def generate_document():
    keyword = keyword_entry.get()
    website_url = url_entry.get()

    if not keyword or not website_url:
        messagebox.showerror("Error", "Both keyword and website URL are required.")
        return

    scraped_data = scrape_website(keyword, website_url)

    if scraped_data:
        create_word_document(keyword, scraped_data)

# Create the main window
root = tk.Tk()
root.title("Data Collection and Document Generation")

# Create and place input fields and labels
keyword_label = tk.Label(root, text="Keyword:")
keyword_label.pack()
keyword_entry = tk.Entry(root)
keyword_entry.pack()

url_label = tk.Label(root, text="Website URL:")
url_label.pack()
url_entry = tk.Entry(root)
url_entry.pack()

# Create and place the "Generate Document" button
generate_button = tk.Button(root, text="Generate Document", command=generate_document)
generate_button.pack()

# Start the GUI event loop
root.mainloop()
